"""
rag_service.py — 医疗专业 RAG 知识库服务（适配层）

保留全部原有 API 接口（RAGService 类的方法签名完全不变），
内部实现委托给 rag_engine.RAGEngine（LangChain + Qdrant）。

文档元数据仍存 MySQL，向量数据迁至 Qdrant。
"""

from __future__ import annotations

import json
import re
from typing import List, Optional

from loguru import logger
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from models import KnowledgeChunk, KnowledgeDocument
from rag_engine import get_rag_engine
from redis_cache import get_cache


def _clean_text(text: str) -> str:
    """清洗文本（保留与原版兼容）。"""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{3,}", " ", text)
    return text.strip()


class RAGService:
    """医疗知识库 RAG 服务（适配层）。

    内部使用 RAGEngine（LangChain + Qdrant）完成文档解析/分块/向量化/检索，
    MySQL 仅存储文档元数据（标题、文件名等）。
    """

    def __init__(self, db: AsyncSession):
        self.db = db
        self._engine = get_rag_engine()

    async def upload_document(
        self,
        content: bytes,
        filename: str,
        user_id: int,
    ) -> int:
        """上传并解析文档，返回文档 ID。"""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

        # 解析文件内容（用于计算 chunk_count）
        if ext == "pdf":
            text = self._parse_pdf(content)
        elif ext in ("docx", "doc"):
            text = self._parse_docx(content)
        else:
            text = content.decode("utf-8", errors="replace")

        text = _clean_text(text)
        if not text or text.startswith("["):
            raise ValueError(f"文件解析失败或内容为空: {text}")

        # 创建文档元数据记录
        title = filename.rsplit(".", 1)[0] if "." in filename else filename
        doc = KnowledgeDocument(
            filename=filename,
            title=title,
            doc_type=ext,
            uploaded_by=user_id,
            chunk_count=0,
        )
        self.db.add(doc)
        await self.db.flush()

        # 委托 RAGEngine：文档解析 → 分块 → 向量化 → Qdrant
        try:
            chunk_count = self._engine.add_document(
                content=content,
                filename=filename,
                document_id=doc.id,
                document_title=title,
            )
        except Exception as e:
            logger.error(f"RAGEngine 入库失败: {e}")
            # 回滚文档元数据
            await self.db.rollback()
            raise ValueError(f"文档入库失败: {e}")

        doc.chunk_count = chunk_count
        await self.db.commit()

        # 知识库更新 → 清除相关缓存
        try:
            cache = get_cache()
            await cache.invalidate_rag_cache()
        except Exception:
            pass

        return doc.id

    async def search(self, query: str, top_k: int = 5) -> List[dict]:
        """语义检索优先走 Qdrant，Redis 缓存加速。"""
        from redis_cache import get_cache

        cache = get_cache()

        # 1. 尝试 Redis 缓存
        try:
            cached = await cache.get_rag_cache(query)
            if cached:
                logger.debug(f"RAG 缓存命中: query={query[:30]}...")
                return cached
        except Exception:
            pass

        # 2. Qdrant 向量检索
        rag_results = self._engine.search(query, top_k)

        # 3. 写入缓存
        if rag_results:
            try:
                await cache.set_rag_cache(query, rag_results, ttl=3600)
            except Exception:
                pass

        return rag_results

    async def delete_document(self, doc_id: int):
        """删除文档（软删除 + 清除 Qdrant 向量 + 清除缓存）。"""
        doc = await self.db.get(KnowledgeDocument, doc_id)
        if not doc:
            return

        doc.is_active = False

        # 清除 MySQL chunks
        await self.db.execute(
            delete(KnowledgeChunk).where(KnowledgeChunk.document_id == doc_id)
        )

        await self.db.commit()

        # 清除 Qdrant 向量
        try:
            self._engine.delete_document(doc_id)
        except Exception as e:
            logger.warning(f"Qdrant 删除失败（非致命）: {e}")

        # 清除缓存
        try:
            cache = get_cache()
            await cache.invalidate_rag_cache()
        except Exception:
            pass

    async def list_documents(self) -> List[KnowledgeDocument]:
        """列出所有活跃文档。"""
        result = await self.db.execute(
            select(KnowledgeDocument)
            .where(KnowledgeDocument.is_active == True)
            .order_by(KnowledgeDocument.created_at.desc())
        )
        return list(result.scalars().all())

    async def cleanup_stale(self) -> int:
        """清理重复内容，返回移除数量。"""
        result = await self.db.execute(
            select(KnowledgeChunk.content, KnowledgeChunk.id)
            .join(KnowledgeDocument, KnowledgeChunk.document_id == KnowledgeDocument.id)
            .where(KnowledgeDocument.is_active == True)
        )
        rows = result.all()

        seen: dict[str, int] = {}
        dup_ids: List[int] = []
        for content, chunk_id in rows:
            key = content[:100] if content else ""
            if key in seen:
                dup_ids.append(chunk_id)
            else:
                seen[key] = chunk_id

        if dup_ids:
            await self.db.execute(
                delete(KnowledgeChunk).where(KnowledgeChunk.id.in_(dup_ids))
            )
            await self.db.commit()

        return len(dup_ids)

    # ── 保留原有的文档解析方法（供 fallback 使用） ─────────

    def _parse_pdf(self, content: bytes) -> str:
        """解析 PDF 文件内容（保留与原版兼容）。"""
        try:
            from PyPDF2 import PdfReader
            import io

            reader = PdfReader(io.BytesIO(content))
            texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            return "\n\n".join(texts)
        except ImportError:
            return "[PDF解析失败: PyPDF2未安装]"
        except Exception as e:
            logger.error(f"PDF解析错误: {e}")
            return f"[PDF解析失败: {e}]"

    def _parse_docx(self, content: bytes) -> str:
        """解析 DOCX 文件内容（保留与原版兼容）。"""
        try:
            from docx import Document
            import io

            doc = Document(io.BytesIO(content))
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(texts)
        except ImportError:
            return "[DOCX解析失败: python-docx未安装]"
        except Exception as e:
            logger.error(f"DOCX解析错误: {e}")
            return f"[DOCX解析失败: {e}]"


# 工厂函数
def get_rag_service(db: AsyncSession) -> RAGService:
    return RAGService(db)
